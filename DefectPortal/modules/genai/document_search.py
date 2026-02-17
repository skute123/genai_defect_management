"""
Document Search Module
Searches knowledge base documents for relevant information.
"""

import logging
import os
from typing import List, Dict, Any, Optional
from pathlib import Path

logger = logging.getLogger(__name__)

class DocumentSearch:
    """
    Service for searching knowledge base documents.
    Processes and indexes documents, then searches for relevant content.
    """
    
    def __init__(self, embedding_service, vector_store, documents_path: str = None):
        """
        Initialize the document search service.
        
        Args:
            embedding_service: EmbeddingService instance.
            vector_store: VectorStore instance.
            documents_path: Path to knowledge base documents folder.
        """
        self.embedding_service = embedding_service
        self.vector_store = vector_store
        
        if documents_path is None:
            base_path = Path(__file__).parent.parent.parent
            documents_path = str(base_path / "knowledge_base" / "documents")
        
        self.documents_path = documents_path
        self.supported_extensions = ['.txt', '.md', '.pdf', '.docx']
    
    def load_and_index_documents(self, force_reindex: bool = False):
        """
        Load all documents from the knowledge base and index them.
        Skips re-indexing if documents are already cached (unless force_reindex=True).
        
        Args:
            force_reindex: Force re-indexing even if cache exists.
        """
        # Check if already indexed (use cache)
        if not force_reindex:
            stats = self.vector_store.get_collection_stats()
            cached_count = stats.get('document_count', 0)
            if cached_count > 0:
                logger.info(f"Using cached document embeddings ({cached_count} chunks already indexed)")
                return
        
        if not os.path.exists(self.documents_path):
            logger.warning(f"Documents path does not exist: {self.documents_path}")
            os.makedirs(self.documents_path, exist_ok=True)
            return
        
        logger.info(f"Loading documents from: {self.documents_path}")
        
        all_chunks = []
        
        # Walk through all files
        for root, dirs, files in os.walk(self.documents_path):
            for file in files:
                filepath = os.path.join(root, file)
                ext = os.path.splitext(file)[1].lower()
                
                if ext in self.supported_extensions:
                    chunks = self._process_file(filepath)
                    all_chunks.extend(chunks)
        
        if not all_chunks:
            logger.warning("No document chunks found to index")
            return
        
        # Clear existing documents before re-indexing
        stats = self.vector_store.get_collection_stats()
        if stats.get('document_count', 0) > 0:
            self.vector_store.clear_documents()
        
        # Generate embeddings
        logger.info(f"Generating embeddings for {len(all_chunks)} document chunks...")
        texts = [chunk['content'] for chunk in all_chunks]
        embeddings = self.embedding_service.generate_embeddings(texts)
        
        # Store in vector database
        self.vector_store.add_documents(all_chunks, embeddings)
        logger.info(f"Successfully indexed {len(all_chunks)} document chunks")
    
    def _process_file(self, filepath: str) -> List[Dict[str, Any]]:
        """
        Process a single file and split into chunks.
        
        Args:
            filepath: Path to the file.
            
        Returns:
            List of chunk dictionaries.
        """
        chunks = []
        filename = os.path.basename(filepath)
        ext = os.path.splitext(filename)[1].lower()
        
        try:
            content = ""
            
            if ext in ['.txt', '.md']:
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            
            elif ext == '.pdf':
                content = self._extract_pdf_text(filepath)
            
            elif ext == '.docx':
                content = self._extract_docx_text(filepath)
            
            if content:
                # Split into chunks
                file_chunks = self._split_text(content, chunk_size=500, overlap=50)
                
                for i, chunk_text in enumerate(file_chunks):
                    chunks.append({
                        'id': f"{filename}_{i}",
                        'content': chunk_text,
                        'filename': filename,
                        'filepath': filepath,
                        'section': self._extract_section(chunk_text),
                        'chunk_index': i
                    })
            
            logger.info(f"Processed {filename}: {len(chunks)} chunks")
            
        except Exception as e:
            logger.error(f"Failed to process {filepath}: {e}")
        
        return chunks
    
    def _extract_pdf_text(self, filepath: str) -> str:
        """Extract text from PDF file."""
        try:
            import pypdf
            text_parts = []
            with open(filepath, 'rb') as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    text_parts.append(page.extract_text() or '')
            return "\n".join(text_parts)
        except ImportError:
            logger.warning("pypdf not installed, skipping PDF files")
            return ""
        except Exception as e:
            logger.error(f"Failed to extract PDF text: {e}")
            return ""
    
    def _extract_docx_text(self, filepath: str) -> str:
        """Extract text from DOCX file (paragraphs and tables)."""
        try:
            from docx import Document
            doc = Document(filepath)
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        parts.append(row_text)
            return "\n".join(parts)
        except ImportError:
            logger.warning("python-docx not installed, skipping DOCX files")
            return ""
        except Exception as e:
            logger.error(f"Failed to extract DOCX text: {e}")
            return ""
    
    def _split_text(self, text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Text to split.
            chunk_size: Target size of each chunk in words.
            overlap: Number of overlapping words between chunks.
            
        Returns:
            List of text chunks.
        """
        words = text.split()
        chunks = []
        
        if len(words) <= chunk_size:
            return [text] if text.strip() else []
        
        start = 0
        while start < len(words):
            end = start + chunk_size
            chunk = " ".join(words[start:end])
            if chunk.strip():
                chunks.append(chunk)
            start = end - overlap
        
        return chunks
    
    def _extract_section(self, text: str) -> str:
        """
        Try to extract a section header from the text chunk.
        
        Args:
            text: Text chunk.
            
        Returns:
            Section header if found, empty string otherwise.
        """
        lines = text.split('\n')
        for line in lines[:3]:  # Check first 3 lines
            line = line.strip()
            # Look for header patterns
            if line and len(line) < 100:
                if line.startswith('#') or line.isupper() or ':' in line:
                    return line.replace('#', '').strip()
        return ""
    
    def _deduplicate_results_by_document(
        self,
        results: List[Dict[str, Any]],
        n_results: int
    ) -> List[Dict[str, Any]]:
        """
        Keep at most one chunk per document (by filename) so the same file
        doesn't appear multiple times (e.g. three chunks from one PDF).
        Preserves order; keeps the first (best) chunk per document.
        """
        seen = set()
        unique = []
        for r in results:
            if len(unique) >= n_results:
                break
            meta = r.get('metadata', {})
            filename = meta.get('filename') or meta.get('filepath') or r.get('id', '')
            key = (filename,) if isinstance(filename, str) else tuple(filename) if filename else (r.get('id'),)
            if key not in seen:
                seen.add(key)
                unique.append(r)
        return unique

    def search(
        self,
        query: str,
        n_results: int = 3,
        min_similarity: float = 0.22
    ) -> List[Dict[str, Any]]:
        """
        Search for relevant documents (semantic + keyword fallback for error-style queries).
        
        Uses a lower default min_similarity so technical/error queries (e.g. KIAS-SetMarketingPermissions)
        still return related docs. If semantic search returns nothing, falls back to keyword
        match so documents that contain the query terms are shown.
        
        Args:
            query: Search query text.
            n_results: Maximum number of results.
            min_similarity: Minimum similarity threshold (0–1). Default 0.22 so error messages still match.
            
        Returns:
            List of relevant document chunks.
        """
        if not query or not query.strip():
            return []
        
        # Semantic search (lower threshold so "An error was encountered while invoking KIAS-SetMarketingPermissions" can match)
        query_embedding = self.embedding_service.generate_embedding(query)
        results = self.vector_store.search_documents(
            query_embedding,
            n_results=n_results,
            min_similarity=min_similarity
        )
        
        # If no semantic hits (e.g. long error message), use keyword fallback so docs containing
        # terms like KIAS, SetMarketingPermissions still appear in Related Knowledge Documents
        if not results:
            results = self.vector_store.search_documents_by_keywords(query, n_results=n_results)
            if results:
                logger.info("Document search used keyword fallback for query terms")
        
        # Deduplicate by document (filename) so we don't show the same document 3 times
        return self._deduplicate_results_by_document(results, n_results)
    
    def search_by_defect(
        self,
        defect: Dict[str, Any],
        n_results: int = 3
    ) -> List[Dict[str, Any]]:
        """
        Search for documents relevant to a specific defect.
        
        Args:
            defect: Defect dictionary.
            n_results: Maximum number of results.
            
        Returns:
            List of relevant document chunks.
        """
        # Create search query from defect fields
        query_parts = [
            defect.get('Summary', ''),
            defect.get('Description', ''),
            defect.get('OSF-System', ''),
            defect.get('Custom field (OSF-Fix Description)', '')
        ]
        
        query = " ".join([str(p) for p in query_parts if p and str(p).lower() != 'nan'])
        
        return self.search(query, n_results=n_results)
    
    def is_indexed(self) -> bool:
        """Check if documents have been indexed."""
        stats = self.vector_store.get_collection_stats()
        return stats.get('document_count', 0) > 0
